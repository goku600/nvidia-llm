import asyncio
import file_modifier
import sys

async def run_omni_test(test_name: str, doc_code: str):
    print(f"\n[{test_name.upper()}]")
    print("-" * 50)
    
    print(f"-> Compiling Sandbox Code ({len(doc_code)} chars)...")
    try:
        output_bytes, output_filename, error = await asyncio.to_thread(
            file_modifier.execute_python_code, doc_code, b""
        )
    except Exception as e:
        print(f"❌ FAIL: Critical Sandbox Crash: {e}")
        return False
        
    if error:
        print(f"❌ FAIL: Python Execution Error:\n{error}")
        return False
        
    if output_bytes:
        print(f"[SUCCESS] Generated '{output_filename}' ({len(output_bytes)} bytes)")
        with open(f"test_omni_{output_filename}", "wb") as f:
            f.write(output_bytes)
    else:
        print("[SUCCESS] Code executed silently with no file returned.")
        
    return True

async def main():
    print("==========================================")
    print("NVIDIA OMNI-MODAL SANDBOX INTEGRATION TEST")
    print("==========================================")
    
    word_test = """
import docx
doc = docx.Document()
doc.add_heading('QuantumLeap Executive Summary', 0)
doc.add_paragraph('This is a test document generated dynamically by the sandbox avoiding token ceilings.')
table = doc.add_table(rows=2, cols=2)
table.cell(0, 0).text = 'Plan'
table.cell(0, 1).text = 'Price'
output_filename = 'quantumleap_proposal.docx'
import io
doc.save(output_buffer)
"""

    scrape_test = """
import requests
from bs4 import BeautifulSoup
import pandas as pd

req = requests.get('https://en.wikipedia.org/wiki/Main_Page', headers={'User-Agent': 'Mozilla'})
req.raise_for_status()

soup = BeautifulSoup(req.content, 'html.parser')
links = soup.select('#mp-otd a')[:3]

data = [{'Title': link.text, 'URL': 'https://en.wikipedia.org' + link.get('href', '')} for link in links]
df = pd.DataFrame(data)

df.to_csv(output_buffer, index=False)
output_filename = 'wiki_otd.csv'
"""

    chart_test = """
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt

months = ['Jan', 'Feb', 'Mar', 'Apr', 'May']
sales = np.random.randint(100, 500, size=5)

plt.figure(figsize=(8, 4))
plt.bar(months, sales, color='skyblue')
plt.title('Monthly Sales Data')

plt.savefig(output_buffer, format='png')
output_filename = 'sales_chart.png'
"""

    tests = [
        ("Word Document Generation", word_test.strip()),
        ("Web Scraping & Data Extraction", scrape_test.strip()),
        ("Data Visualization", chart_test.strip())
    ]
    
    passed = 0
    for name, code in tests:
        success = await run_omni_test(name, code)
        if success:
            passed += 1
            
    print("\n" + "="*42)
    print(f"TEST RESULTS: {passed}/{len(tests)} PASSED")
    print("==========================================")
    
    if passed == len(tests):
        sys.exit(0)
    else:
        sys.exit(1)

if __name__ == "__main__":
    asyncio.run(main())

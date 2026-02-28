import io
import asyncio
import logging
import time

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, Message
from telegram.ext import ContextTypes
from telegram.constants import ParseMode, ChatAction

import nvidia_client as ai
import session as sess
import file_modifier
from config import SELECTABLE_MODELS, ALLOWED_USER_IDS, AUTO_DELETE_DOC_AFTER_ANSWER, DOC_TTL_SECONDS

logger = logging.getLogger(__name__)


async def _keep_typing(bot, chat_id: int, stop_event: asyncio.Event):
    """Send typing action every 4 seconds until stop_event is set."""
    while not stop_event.is_set():
        try:
            await bot.send_chat_action(chat_id=chat_id, action=ChatAction.TYPING)
        except Exception:
            pass
        await asyncio.sleep(4)


def _is_allowed(user_id: int) -> bool:
    """Return True if user is allowed. If allowlist is empty, everyone is allowed."""
    if not ALLOWED_USER_IDS:
        return True
    return user_id in ALLOWED_USER_IDS


async def _check_allowed(update: Update) -> bool:
    """Check allowlist and send rejection if not allowed. Returns True if allowed."""
    if not _is_allowed(update.effective_user.id):
        await update.effective_message.reply_text(
            "⛔ Sorry, you are not authorized to use this bot.\n"
            "Contact the bot owner to request access."
        )
        return False
    return True

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

MAX_MSG = 4000

async def _edit_or_send(thinking_msg: Message, update: Update, text: str):
    """
    Edit the thinking_msg with the full reply if it fits in one message.
    If the text is too long, edit with the first chunk and send the rest as
    a single follow-up message (truncated with a note) to avoid choppy multi-message UX.
    """
    if not text or not text.strip():
        # Fallback if the AI returns literally nothing
        text = "(No spoken response)"
        
    try:
        if len(text) <= MAX_MSG:
            try:
                await thinking_msg.edit_text(text, parse_mode=ParseMode.MARKDOWN)
            except Exception as e:
                if "Message is not modified" in str(e): return
                await thinking_msg.edit_text(text) # Fallback to plaintext
        else:
            # Send first MAX_MSG chars, then one follow-up with the remainder
            first = text[:MAX_MSG]
            rest = text[MAX_MSG:]
            try:
                await thinking_msg.edit_text(first, parse_mode=ParseMode.MARKDOWN)
            except Exception:
                await thinking_msg.edit_text(first)
            
            # If rest is also too long, truncate gracefully
            if len(rest) > MAX_MSG:
                rest = rest[:MAX_MSG - 100] + "\n\n_[Response truncated — ask me to continue]_"
            
            try:
                await update.effective_message.reply_text(rest, parse_mode=ParseMode.MARKDOWN)
            except Exception:
                await update.effective_message.reply_text(rest)
    except Exception as e:
        logger.error(f"Failed to edit message entirely in _edit_or_send: {e}")


def _is_modification_request(text: str) -> bool:
    """Detect if the user wants to modify/edit the file and get it back."""
    strong_keywords = [
        "modify the file", "edit the file", "change the file", "update the file",
        "modify this file", "edit this file",
        "send me the file", "send the file back", "send back", "send me back",
        "download the file", "export the file", "give me the file",
        "save the file", "save and send", "save it",
        "add a row", "add a column", "add a sheet", "add a total", "add a sum",
        "remove the row", "remove the column", "delete the row", "delete the column",
        "remove duplicates", "delete duplicates",
        "sort the", "sort by", "filter the", "filter by",
        "rename the column", "rename column",
        "reformat the", "convert to csv", "convert to excel", "convert to pdf",
        "convert the file", "convert this to",
        "insert a row", "insert a column",
        "replace all", "find and replace",
        "clean the data", "clean up the",
        "merge the sheets", "split the file",
        "format the", "fix the file",
        "make changes", "apply changes",
        "update the data", "update the row", "update the column",
        "fill in the", "fill the",
        "append to", "prepend to",
    ]
    text_lower = text.lower()
    return any(kw in text_lower for kw in strong_keywords)


def _is_create_file_request(text: str) -> tuple[bool, str]:
    """
    Detect if the user wants to create a new file from scratch.
    Returns (is_create_request, detected_extension).
    """
    text_lower = text.lower()

    # Must have a creation verb
    create_verbs = [
        "create", "generate", "make", "write", "produce", "build",
        "give me a", "give me the", "send me a", "send me the",
        "i want a", "i need a", "can you make", "can you create",
        "can you generate", "can you write",
    ]
    has_verb = any(v in text_lower for v in create_verbs)
    if not has_verb:
        return False, ""

    # Map keywords to file extensions
    file_type_map = [
        (["text file", ".txt", "txt file", "plain text file"], "txt"),
        (["excel file", ".xlsx", "spreadsheet", "excel sheet", "xlsx"], "xlsx"),
        (["csv file", ".csv", "csv"], "csv"),
        (["word file", "word doc", ".docx", "docx", "word document"], "docx"),
        (["json file", ".json", "json"], "json"),
        (["python script", "python file", ".py", "py script", "python program"], "py"),
        (["markdown file", ".md", "md file", "markdown"], "md"),
        (["html file", ".html", "html page", "webpage"], "html"),
        (["javascript file", ".js", "js file", "javascript"], "js"),
        (["yaml file", ".yaml", ".yml", "yaml"], "yaml"),
        (["xml file", ".xml", "xml"], "xml"),
        (["pdf", ".pdf"], "pdf"),
    ]

    for keywords, ext in file_type_map:
        if any(kw in text_lower for kw in keywords):
            return True, ext

    # Generic "file" or "document" without specific type → default to txt
    if any(w in text_lower for w in ["a file", "a document", "the file", "the document"]):
        return True, "txt"

    return False, ""


def _is_image_followup(text: str) -> bool:
    """Detect if the user wants to refer to the recently uploaded image."""
    followup_keywords = [
        "image", "photo", "picture", "pic", "it", "this", "translate", "read",
        "extract", "what does", "can you see", "in there", "on there",
        "that", "the man", "the cat", "the dog", "the document", "the written",
        "saying", "text"
    ]
    text_lower = text.lower()
    # It's likely to be a follow up if it contains these specific words
    if any(kw in text_lower for kw in followup_keywords):
        return True
    return False


# ---------------------------------------------------------------------------
# Commands
# ---------------------------------------------------------------------------

async def cmd_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not await _check_allowed(update): return
    user = update.effective_user
    sess.clear_session(user.id)
    await update.message.reply_text(
        f"👋 Hello, {user.first_name}! I'm your NVIDIA AI-powered assistant.\n\n"
        "Just send me anything and we can chat normally:\n\n"
        "💬 *Chat* — Ask questions, brainstorm, or converse\n"
        "👨‍💻 *Code* — Write, debug & explain code\n"
        "📄 *Files* — Send any PDF, Word, Excel, CSV, TXT... I'll read it AND edit it for you!\n"
        "🖼️ *Images* — Send any photo and I'll analyze it!\n\n"
        "It's all one continuous conversation. 🚀",
        parse_mode=ParseMode.MARKDOWN,
    )


async def cmd_clear(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not await _check_allowed(update): return
    user_id = update.effective_user.id
    sess.clear_session(user_id)
    await update.message.reply_text(
        "🗑️ All data cleared — conversation history, documents, and session reset!"
    )


async def cmd_privacy(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not await _check_allowed(update): return
    user_id = update.effective_user.id
    info = sess.get_privacy_info(user_id)

    ttl_mins = DOC_TTL_SECONDS // 60
    doc_status = "None"
    if info["has_doc"]:
        expires_mins = (info["doc_expires_in"] or 0) // 60
        doc_status = f"✅ *{info['doc_name']}* (expires in ~{expires_mins} min)"

    history_note = (
        f"{info['history_count']} messages in memory"
        if info["history_count"] > 0
        else "No messages stored"
    )

    await update.message.reply_text(
        "🔒 *Your Privacy & Data Summary*\n\n"
        f"*AI model:* `{info['model']}`\n"
        f"*Conversation:* {history_note}\n"
        f"*Document info context:* {doc_status}\n\n"
        "📋 *What we store (in RAM only):*\n"
        "• Your continuous conversation history (last 20 messages)\n"
        "• Uploaded document text (auto-deleted after "
        f"{ttl_mins} min)\n"
        "• Your selected model\n\n"
        "📋 *What we DON'T store:*\n"
        "• Files on disk (everything is in RAM)\n"
        "• Your Telegram messages permanently\n"
        "• Any data after bot restarts\n\n"
        "⚠️ *Note:* Text is sent to NVIDIA's API for processing.\n\n"
        "🗑️ Use /clear to delete all your data immediately.",
        parse_mode=ParseMode.MARKDOWN,
    )


async def cmd_help(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not await _check_allowed(update): return
    await update.message.reply_text(
        "🤖 *NVIDIA AI Telegram Bot — Help*\n\n"
        "🪄 *Omni-modal Assistant:* It's all one conversation!\n"
        "• Send a *photo* → it understands the image context\n"
        "• Send a *file* → it reads the document into your chat history\n"
        "• Ask to *create or modify* a file → I will write code and send the file to you!\n\n"
        "*Commands:*\n"
        "/start — Welcome message\n"
        "/model — Switch the AI model\n"
        "/clear — Clear all your data\n"
        "/privacy — View & manage your stored data\n"
        "/help — Show this help message\n\n"
        "*Supported file types:*\n"
        "PDF, Word (.docx), Excel (.xlsx), CSV, JSON, TXT, MD, PY and more\n\n"
        "*Tips:*\n"
        "• Bot remembers your last 20 messages\n"
        "• You can ask me to generate a random PDF or process an Excel sheet, and I'll send the file straight back.",
        parse_mode=ParseMode.MARKDOWN,
    )


def _model_keyboard():
    buttons = []
    for key, (model_id, label) in SELECTABLE_MODELS.items():
        buttons.append([InlineKeyboardButton(label, callback_data=f"model_{key}")])
    return InlineKeyboardMarkup(buttons)


async def cmd_model(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    current_model = sess.get_model(user_id)
    # Find current model label
    current_label = current_model
    for key, (model_id, label) in SELECTABLE_MODELS.items():
        if model_id == current_model:
            current_label = label
            break
    await update.message.reply_text(
        f"🧠 *Current model:* `{current_label}`\n\n"
        "Choose a model to switch to:\n"
        "_(Note: Image Analysis always uses the vision model regardless of this setting)_",
        parse_mode=ParseMode.MARKDOWN,
        reply_markup=_model_keyboard(),
    )


# ---------------------------------------------------------------------------
# Callback
# ---------------------------------------------------------------------------
    query = update.callback_query
    await query.answer()
    user_id = query.from_user.id
    key = query.data.replace("model_", "")

    if key not in SELECTABLE_MODELS:
        return

    model_id, label = SELECTABLE_MODELS[key]
    sess.set_model(user_id, model_id)

    await query.edit_message_text(
        f"✅ Switched to *{label}*\n\n"
        f"`{model_id}`\n\n"
        "Conversation history cleared. Start chatting!",
        parse_mode=ParseMode.MARKDOWN,
    )


# ---------------------------------------------------------------------------
# Text messages
# ---------------------------------------------------------------------------

async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not await _check_allowed(update): return
    user_id = update.effective_user.id
    user_text = update.message.text.strip()

    if not user_text:
        return

    # Send immediate feedback
    thinking_msg: Message = await update.message.reply_text("⏳ Thinking...")

    # Start persistent typing indicator in background
    stop_typing = asyncio.Event()
    typing_task = asyncio.create_task(
        _keep_typing(context.bot, update.effective_chat.id, stop_typing)
    )

    sess.add_message(user_id, "user", user_text)
    history = sess.get_history(user_id)
    model = sess.get_model(user_id)
    
    # Check if there is an active image in the session for vision follow-ups
    image_bytes, image_mime = sess.get_image_file(user_id)
    
    try:
        # Consume the streaming generator inside a loop to support tool use (like Web Search)
        output_file = None
        final_reply = ""
        last_edit_time = 0
        
        while True:
            full_reply = ""
            
            # Route to vision model if image exists in context AND user's text seems related
            # Otherwise use standard text model
            if image_bytes and _is_image_followup(user_text):
                # remove the last user message from history, as it's passed as prompt
                chat_history = history[:-1]
                reply_generator = ai.image_analysis(image_bytes, image_mime or "image/jpeg", user_text, chat_history)
            else:
                reply_generator = ai.chat(history, model=model)
            
            async for chunk in reply_generator:
                full_reply += chunk
                
                # Update Telegram message every 1.5 seconds to avoid rate limits
                current_time = time.time()
                if current_time - last_edit_time > 1.5:
                    # Truncate for streaming preview if it gets too long
                    preview_text = full_reply if len(full_reply) < MAX_MSG else full_reply[:MAX_MSG] + "..."
                    try:
                        await thinking_msg.edit_text(preview_text + " ⏳")
                        last_edit_time = current_time
                    except Exception:
                        pass
            
            # 1. Check for Web Search Tool
            search_start = full_reply.find("[WEB_SEARCH]")
            search_end = full_reply.find("[/WEB_SEARCH]")
            
            if search_start != -1 and search_end != -1 and search_end > search_start:
                query = full_reply[search_start + len("[WEB_SEARCH]"):search_end].strip()
                await thinking_msg.edit_text(f"🔍 Searching the web for: `{query}`...", parse_mode=ParseMode.MARKDOWN)
                
                import tools
                search_results = await asyncio.to_thread(tools.search_web, query)
                
                # We save the AI's partial thought
                clean_thought = full_reply[:search_start].strip()
                if clean_thought:
                    sess.add_message(user_id, "assistant", clean_thought)
                
                # Inject the web results into history as systemic context
                context_msg = f"Web search results for '{query}':\n\n{search_results}\n\nNow provide your final answer based on these results."
                sess.add_message(user_id, "system", context_msg)
                
                # Re-fetch history and loop again!
                history = sess.get_history(user_id)
                continue

            # 2. Check for File Execution Tool
            exec_start = full_reply.find("[PYTHON_EXEC]")
            exec_end = full_reply.find("[/PYTHON_EXEC]")
            
            if exec_start != -1 and exec_end != -1 and exec_end > exec_start:
                await thinking_msg.edit_text("⚙️ Generating your file, please wait...")
                code = full_reply[exec_start + len("[PYTHON_EXEC]"):exec_end].strip()
                
                # Fetch explicitly uploaded file bytes if any exist in session
                file_bytes, _, _ = sess.get_doc_file(user_id)
                input_bytes = file_bytes if file_bytes else b""
                
                output_bytes, output_filename, error = await asyncio.to_thread(
                    file_modifier.execute_python_code, code, input_bytes
                )
                
                # Clean the reply to remove the python block
                clean_reply = full_reply[:exec_start].strip() + "\n" + full_reply[exec_end + len("[/PYTHON_EXEC]"):].strip()
                final_reply = clean_reply.strip() or "✅ Task completed."
                
                if error:
                    final_reply += f"\n\n⚠️ **Execution Error:**\n```\n{error[:500]}\n```"
                elif output_bytes:
                    output_file = (output_bytes, output_filename)
                else:
                    final_reply += "\n\n⚠️ No file was generated by the executed code."
            else:
                # Normal terminal reply
                final_reply = full_reply
            
            # If we get here (we didn't trigger a web search 'continue'), we must break the loop
            break

    except Exception as e:
        logger.error(f"NVIDIA API error: {e}")
        stop_typing.set()
        typing_task.cancel()
        final_reply = f"⚠️ Sorry, I ran into an error talking to the AI:\n```\n{e}\n```"
        await _edit_or_send(thinking_msg, update, final_reply)
        return

    finally:
        stop_typing.set()
        typing_task.cancel()

    # Log assistant response to history
    sess.add_message(user_id, "assistant", final_reply)
    await _edit_or_send(thinking_msg, update, final_reply)
    
    # If a file was generated, send it as a follow-up
    if output_file:
        out_b, out_name = output_file
        await update.message.reply_document(
            document=io.BytesIO(out_b),
            filename=out_name,
            caption=f"✅ Here's your file: *{out_name}*",
            parse_mode=ParseMode.MARKDOWN,
        )


# ---------------------------------------------------------------------------
# Photo messages
# ---------------------------------------------------------------------------

async def handle_photo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not await _check_allowed(update): return
    user_id = update.effective_user.id

    thinking_msg: Message = await update.message.reply_text("⏳ Analyzing image...")

    stop_typing = asyncio.Event()
    typing_task = asyncio.create_task(
        _keep_typing(context.bot, update.effective_chat.id, stop_typing)
    )

    photo = update.message.photo[-1]
    caption = update.message.caption or "Describe this image in detail."

    try:
        from config import VISION_MODEL
        import base64
        
        file = await context.bot.get_file(photo.file_id)
        file_bytes = bytes(await file.download_as_bytearray())
        image_b64 = base64.b64encode(file_bytes).decode('utf-8')

        history = sess.get_history(user_id)
        
        # Save the image bytes to session so subsequent text messages can still "see" it
        sess.set_image_file(user_id, file_bytes, "image/jpeg")
        
        # We don't save the massive base64 string to the session history to save RAM.
        # We just log that an image was sent for context.
        sess.add_message(user_id, "user", f"[User sent an image: {caption}]")
        
        # Process via the dedicated image_analysis function to avoid system prompt overriding the vision constraints
        reply_generator = ai.image_analysis(file_bytes, "image/jpeg", caption, history)
        
        reply = ""
        async for chunk in reply_generator:
            reply += chunk
            # Optional: Add streaming UI update here if needed
            # await thinking_msg.edit_text(reply + " ⏳")

        sess.add_message(user_id, "assistant", reply)
        await _edit_or_send(thinking_msg, update, reply)

    except Exception as e:
        logger.error(f"Image analysis error: {e}")
        await thinking_msg.edit_text(
            f"⚠️ Error analyzing the image:\n`{e}`",
            parse_mode=ParseMode.MARKDOWN,
        )
    finally:
        stop_typing.set()
        typing_task.cancel()


# ---------------------------------------------------------------------------
# Document messages
# ---------------------------------------------------------------------------

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not await _check_allowed(update): return
    user_id = update.effective_user.id

    thinking_msg: Message = await update.message.reply_text("⏳ Processing document...")

    stop_typing = asyncio.Event()
    typing_task = asyncio.create_task(
        _keep_typing(context.bot, update.effective_chat.id, stop_typing)
    )

    doc = update.message.document
    file_name = doc.file_name or "document"
    mime = doc.mime_type or ""

    try:
        file = await context.bot.get_file(doc.file_id)
        file_bytes = await file.download_as_bytearray()

        # Extract text based on file type
        text = _extract_text(bytes(file_bytes), file_name, mime)

        if not text.strip():
            await thinking_msg.edit_text(
                "⚠️ I couldn't extract any text from that file. "
                "Please try a plain text file (.txt, .md, .py, .csv, .json, etc.)."
            )
            return

        # Limit document size (approx 100k chars to fit in context)
        MAX_DOC_CHARS = 100_000
        if len(text) > MAX_DOC_CHARS:
            text = text[:MAX_DOC_CHARS]
            await update.message.reply_text(
                f"⚠️ Document is large — I'll use the first {MAX_DOC_CHARS:,} characters."
            )

        # Inject into unified history
        doc_entry = f"The user uploaded a document named '{file_name}'. Here is its content:\n\n{text}"
        sess.add_message(user_id, "system", doc_entry)
        
        # We also store the bytes temporarily in case a subsequent script wants to manipulate the raw file
        sess.set_doc_text(user_id, text, filename=file_name, file_bytes=bytes(file_bytes), mime=mime)

        word_count = len(text.split())
        await thinking_msg.edit_text(
            f"✅ *{file_name}* loaded into context!\n"
            f"📊 ~{word_count:,} words extracted.\n\n"
            "I remember this document now. What would you like to do with it?",
            parse_mode=ParseMode.MARKDOWN,
        )

    except Exception as e:
        logger.error(f"Document processing error: {e}")
        await thinking_msg.edit_text(
            f"⚠️ Error processing the document:\n`{e}`",
            parse_mode=ParseMode.MARKDOWN,
        )
    finally:
        stop_typing.set()
        typing_task.cancel()


def _extract_text(data: bytes, filename: str, mime: str) -> str:
    """Extract plain text from file bytes. Supports txt, pdf, docx, xlsx, csv, json, md, py, etc."""
    fname_lower = filename.lower()

    # PDF
    if fname_lower.endswith(".pdf") or "pdf" in mime:
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(data))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(p for p in pages if p.strip())
        except Exception as e:
            return f"[PDF extraction error: {e}]"

    # Word (.docx)
    if fname_lower.endswith(".docx") or "wordprocessingml" in mime or "msword" in mime:
        try:
            import docx
            doc = docx.Document(io.BytesIO(data))
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    parts.append("\t".join(cell.text for cell in row.cells))
            return "\n".join(parts)
        except Exception as e:
            return f"[Word extraction error: {e}]"

    # Excel (.xlsx, .xls)
    if fname_lower.endswith((".xlsx", ".xls")) or "spreadsheet" in mime or "excel" in mime:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            parts = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                parts.append(f"=== Sheet: {sheet_name} ===")
                for row in ws.iter_rows(values_only=True):
                    row_str = "\t".join(str(c) if c is not None else "" for c in row)
                    if row_str.strip():
                        parts.append(row_str)
            return "\n".join(parts)
        except Exception as e:
            return f"[Excel extraction error: {e}]"

    # CSV — decode as text
    if fname_lower.endswith(".csv") or "csv" in mime:
        try:
            return data.decode("utf-8")
        except UnicodeDecodeError:
            return data.decode("latin-1", errors="replace")

    # JSON — decode and pretty print
    if fname_lower.endswith(".json") or "json" in mime:
        try:
            import json
            parsed = json.loads(data.decode("utf-8"))
            return json.dumps(parsed, indent=2, ensure_ascii=False)
        except Exception:
            return data.decode("utf-8", errors="replace")

    # Everything else: plain text (py, md, txt, html, yaml, toml, etc.)
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return data.decode("latin-1")
        except Exception:
            return "[Could not extract text from this file type.]"
